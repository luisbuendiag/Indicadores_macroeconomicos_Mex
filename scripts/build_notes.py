"""Genera notas individuales (.docx) por indicador a partir de machotes aprobados.

- Busca el machote en data/source/notas_machote/{CLAVE}_machote.docx.
- Extrae el periodo que cubre el machote y lo compara con el último periodo validado.
- Si coinciden y el indicador está ACTUALIZADO, copia el machote a
  downloads/indicadores/{CLAVE}/{CLAVE}_nota.docx, sin modificar el machote.
- Si no hay machote, el periodo no coincide o el indicador no está actualizado,
  deja el botón NOTA deshabilitado con la causa correspondiente.
"""
from __future__ import annotations

import argparse
import json
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

import lib_data as L
from sources import inegi

ROOT = Path(__file__).resolve().parents[1]
MACHOTES_DIR = ROOT / "data" / "source" / "notas_machote"
NOTES_DIR = ROOT / "downloads" / "indicadores"

_MES = {
    "enero": 1, "febrero": 2, "marzo": 3, "abril": 4, "mayo": 5, "junio": 6,
    "julio": 7, "agosto": 8, "septiembre": 9, "octubre": 10, "noviembre": 11, "diciembre": 12,
    "ene": 1, "feb": 2, "mar": 3, "abr": 4, "may": 5, "jun": 6,
    "jul": 7, "ago": 8, "sep": 9, "oct": 10, "nov": 11, "dic": 12,
}


def _doc_text(doc: Document) -> str:
    """Extrae todo el texto de un documento Word (párrafos + tablas)."""
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_periods_from_text(text: str, trimestral: bool = False) -> list[tuple[int, int]]:
    """Busca periodos mensuales o trimestrales en el texto del machote.

    Devuelve una lista de (año, mes) candidatos.  Para rangos como
    "ene-jun 2026" o "enero-junio 2026" se usa el mes final.

    Cuando trimestral=True, sólo se buscan patrones trimestrales para evitar
    confundir la fecha de publicación con el periodo de referencia.
    """
    found: list[tuple[int, int]] = []
    text = text.lower()

    # Elimina líneas de publicación que suelen contener meses-año
    # y no corresponden al periodo de referencia.
    text = re.sub(r".*boletín de?l? indicador del.*", "", text)
    text = re.sub(r".*próxima publicación.*", "", text)

    if trimestral:
        # Trimestral: 1T-2026, 1T 2026, 1T-26
        for m in re.finditer(r"\b([1-4])t[\s\-]+(\d{2,4})\b", text, re.IGNORECASE):
            try:
                q = int(m.group(1))
                y = int(m.group(2))
                if y < 100:
                    y = 2000 + y
                found.append((y, (q - 1) * 3 + 1))
            except (KeyError, ValueError):
                pass
        return found

    # Mes completo o abreviado + "de" + año
    for m in re.finditer(
        r"\b(enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre|ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)\s+de\s+(\d{4})\b",
        text,
    ):
        try:
            found.append((int(m.group(2)), _MES[m.group(1)]))
        except KeyError:
            pass

    # Rango tipo "ene-jun 2026" o "enero-junio 2026"; usamos el mes final.
    rango_pat = re.compile(
        r"\b(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre|ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)\s*[-–]\s*(enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre|ene|feb|mar|abr|may|jun|jul|ago|sep|oct|nov|dic)\s+(\d{4})\b"
    )
    for m in rango_pat.finditer(text):
        try:
            found.append((int(m.group(2)), _MES[m.group(1)]))
        except KeyError:
            pass

    return found


def _machote_periods(machote_path: Path, trimestral: bool = False) -> list[tuple[int, int]]:
    try:
        doc = Document(str(machote_path))
        text = _doc_text(doc)
    except Exception:
        return []
    return _extract_periods_from_text(text, trimestral=trimestral)


def _current_period_to_ym(period: str | None) -> tuple[int, int] | None:
    if not period:
        return None
    ym = inegi.label_to_ym(period)
    if not ym:
        return None
    y, m = map(int, ym.split("-"))
    return y, m


def _current_period_label(period: str | None) -> str:
    return period or "—"


def _note_period(out_path: Path, trimestral: bool = False) -> tuple[int, int] | None:
    """Lee el periodo de una nota ya generada para decidir si regenerar."""
    if not out_path.exists():
        return None
    try:
        doc = Document(str(out_path))
        periods = _extract_periods_from_text(_doc_text(doc), trimestral=trimestral)
        return periods[-1] if periods else None
    except Exception:
        return None


def _make_nota_metadata(ind: dict, machote_path: Path, out_path: Path, matched: bool) -> dict:
    return {
        "nota_disponible": matched,
        "nota_causa": None if matched else "Periodo del machote no coincide con periodo validado",
        "url_nota_individual": str(out_path.relative_to(ROOT)) if matched and out_path.exists() else None,
        "nota_periodo": ind.get("last_observation"),
        "nota_fecha_generacion": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "nota_estado": "GENERADA" if matched else "SIN_MACHOTE",
        "nota_machote_usado": str(machote_path.relative_to(ROOT)) if matched else None,
    }


def build_notes(payload: dict, pilot: list[str] | None = None) -> dict[str, Path]:
    keys = pilot if pilot else list(payload["indicators"].keys())
    generated: dict[str, Path] = {}

    NOTES_DIR.mkdir(parents=True, exist_ok=True)
    for key in keys:
        ind = payload["indicators"][key]
        out_dir = NOTES_DIR / key
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"{key}_nota.docx"
        machote_path = MACHOTES_DIR / f"{key}_machote.docx"

        # Reset básico
        for k in ("nota_disponible", "nota_causa", "url_nota_individual", "nota_periodo",
                  "nota_fecha_generacion", "nota_estado", "nota_machote_usado"):
            ind.pop(k, None)

        if not machote_path.exists():
            ind.update(_make_nota_metadata(ind, machote_path, out_path, False))
            ind["nota_causa"] = "Falta machote aprobado"
            ind["nota_estado"] = "SIN_MACHOTE"
            if out_path.exists():
                # No borramos una nota previa válida si aún existe, solo deshabilitamos el botón.
                pass
            continue

        if not ind.get("observations"):
            ind.update(_make_nota_metadata(ind, machote_path, out_path, False))
            ind["nota_causa"] = "Sin observaciones"
            continue

        if ind.get("estado") not in ("ACTUALIZADO",):
            ind.update(_make_nota_metadata(ind, machote_path, out_path, False))
            ind["nota_causa"] = f"Estado del indicador: {ind.get('estado')}. La nota solo se genera para indicadores ACTUALIZADO."
            # Conserva la nota previa si existe, no la sobrescribe.
            if out_path.exists():
                ind["url_nota_individual"] = str(out_path.relative_to(ROOT))
            continue

        current_ym = _current_period_to_ym(ind.get("last_observation"))
        trimestral = "T-" in (ind.get("last_observation") or "")
        machote_periods = _machote_periods(machote_path, trimestral=trimestral)
        existing_ym = _note_period(out_path, trimestral=trimestral)

        if not current_ym:
            ind.update(_make_nota_metadata(ind, machote_path, out_path, False))
            ind["nota_causa"] = "No se pudo determinar el periodo validado"
            continue

        match = current_ym in machote_periods

        if not match:
            # Si no coincide el periodo del machote con el dato validado, NO generamos.
            # Si la nota previa sigue siendo válida para su propio periodo, la conservamos.
            ind.update(_make_nota_metadata(ind, machote_path, out_path, False))
            ind["nota_causa"] = (
                f"Machote cubre {machote_periods} pero el periodo validado es {_current_period_label(ind.get('last_observation'))}. "
                "No se genera nota hasta contar con machote correspondiente."
            )
            if out_path.exists():
                ind["url_nota_individual"] = str(out_path.relative_to(ROOT))
            continue

        # Si ya existe una nota para el mismo periodo, no regenerar innecesariamente.
        if out_path.exists() and existing_ym == current_ym and ind.get("nota_disponible"):
            ind["nota_disponible"] = True
            ind["nota_causa"] = None
            ind["url_nota_individual"] = str(out_path.relative_to(ROOT))
            ind["nota_periodo"] = ind.get("last_observation")
            ind["nota_fecha_generacion"] = datetime.fromtimestamp(out_path.stat().st_mtime, tz=timezone.utc).isoformat(timespec="seconds")
            ind["nota_estado"] = "GENERADA"
            ind["nota_machote_usado"] = str(machote_path.relative_to(ROOT))
            generated[key] = out_path
            continue

        # Generar nota: copiar el machote y registrar metadatos.
        try:
            from shutil import copyfile
            copyfile(machote_path, out_path)
            ind.update(_make_nota_metadata(ind, machote_path, out_path, True))
            ind["nota_estado"] = "GENERADA"
            generated[key] = out_path
        except Exception as e:
            ind.update(_make_nota_metadata(ind, machote_path, out_path, False))
            ind["nota_causa"] = f"Error al generar nota: {e}"

    return generated


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--pilot", action="store_true", help="Solo indicadores con machote")
    args = ap.parse_args()

    payload = L.load_data()
    generated = build_notes(payload)
    (L.DATA_DIR / "indicadores.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2, default=str) + "\n",
        encoding="utf-8",
    )
    if generated:
        print(f"OK: {len(generated)} notas generadas:")
        for key, path in sorted(generated.items()):
            print(f"  - {key}: {path}")
    else:
        print("Aviso: no se generaron notas.")


if __name__ == "__main__":
    main()
